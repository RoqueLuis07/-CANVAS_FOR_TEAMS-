"""Rellena los .docx REALES y autorizados de las Actas (no una
reconstrucción) con los datos cargados, preservando el formato original
intacto — se editan los runs/celdas existentes en una copia del template,
nunca se reescribe el layout. El PDF final sale de convertir ese .docx ya
completado (ver docx_to_pdf.py), así que es *literalmente* el documento
autorizado, no una aproximación visual.
"""
import io
from pathlib import Path

import docx

_TEMPLATES_DIR = Path(__file__).parent.parent / "assets" / "templates"
_TPL_SERVICIO_TECNICO = _TEMPLATES_DIR / "acta_servicio_tecnico.docx"
_TPL_ENTREGA_EQUIPO = _TEMPLATES_DIR / "acta_entrega_equipo.docx"


def _v(val) -> str:
    return str(val).strip() if val else ""


def _set_paragraph_text(paragraph, value: str):
    """Escribe `value` en el primer run del párrafo (creándolo si no
    existe ninguno), preservando el formato de ese run si ya existía."""
    if paragraph.runs:
        paragraph.runs[0].text = value
        for extra in paragraph.runs[1:]:
            extra.text = ""
    else:
        paragraph.add_run(value)


def _set_cell(table, row: int, col: int, value: str):
    """Escribe en una celda VACÍA sin destruir el shading/estilo de la
    celda (que vive en tcPr, no en el texto) — cell.text = ... reemplaza
    el contenido del párrafo pero no toca tcPr."""
    table.cell(row, col).text = value


def _check(paragraph, options: dict[str, bool]):
    """Marca ☑/☐ en un párrafo que tiene un run "☐" seguido de un run con
    el texto de la opción, para cada opción en `options` (texto exacto del
    run -> si va marcada)."""
    runs = paragraph.runs
    for i, r in enumerate(runs):
        if r.text.strip() == "☐":
            # el run siguiente (o el que sigue a uno de espacio) trae la
            # etiqueta de la opción
            label_run = None
            for j in range(i + 1, min(i + 3, len(runs))):
                if runs[j].text.strip():
                    label_run = runs[j]
                    break
            if label_run is None:
                continue
            for opt_text, checked in options.items():
                if label_run.text.strip().startswith(opt_text):
                    if checked:
                        r.text = "☑"
                    break


def generar_docx_servicio_tecnico(a: dict) -> bytes:
    d = docx.Document(str(_TPL_SERVICIO_TECNICO))

    # Tabla 0: N° DE ACTA | FECHA | HORA
    _set_cell(d.tables[0], 1, 0, _v(a.get("numero")))
    _set_cell(d.tables[0], 1, 1, _v(a.get("fecha")))
    _set_cell(d.tables[0], 1, 2, _v(a.get("hora")))

    # Tabla 1: Departamento / Ubicación / Persona que reporta
    _set_cell(d.tables[1], 0, 1, _v(a.get("departamento_area")))
    _set_cell(d.tables[1], 1, 1, _v(a.get("ubicacion_oficina")))
    _set_cell(d.tables[1], 2, 1, _v(a.get("persona_reporta")))

    # Checkboxes de tipo de equipo (párrafo 4)
    tipo = (a.get("tipo_equipo") or "").strip()
    otro_txt = a.get("tipo_equipo_otro") or ""
    _check(d.paragraphs[4], {
        "Computadora": tipo == "Computadora",
        "Impresora": tipo == "Impresora",
        "Teléfono": tipo == "Teléfono",
        "Otro": tipo == "Otro",
    })
    if tipo == "Otro" and otro_txt:
        for r in d.paragraphs[4].runs:
            if r.text.startswith("Otro (especificar):"):
                r.text = f"Otro (especificar): {otro_txt}"
                break

    # Tabla 2: Marca | Modelo | N° de serie
    _set_cell(d.tables[2], 1, 0, _v(a.get("marca")))
    _set_cell(d.tables[2], 1, 1, _v(a.get("modelo")))
    _set_cell(d.tables[2], 1, 2, _v(a.get("nro_serie")))

    # Falla / motivo reportado — párrafo 7 (línea en blanco debajo del label)
    _set_paragraph_text(d.paragraphs[7], _v(a.get("falla_motivo")))

    # Trabajo realizado — párrafo 10
    _set_paragraph_text(d.paragraphs[10], _v(a.get("trabajo_realizado")))

    # Tabla 3: Repuestos/insumos | Técnico responsable
    _set_cell(d.tables[3], 0, 1, _v(a.get("repuestos_insumos")))
    _set_cell(d.tables[3], 1, 1, _v(a.get("tecnico_responsable")))

    # Tabla 4: Firmas (Encargado de TI | Usuario Conforme)
    _fill_firma_cell(d.tables[4].rows[0].cells[0], a.get("encargado_ti_ci"), a.get("encargado_ti_nombre"))
    _fill_firma_cell(d.tables[4].rows[0].cells[1], a.get("usuario_ci"), a.get("usuario_nombre"))

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _fill_firma_cell(cell, ci: str, nombre: str):
    # Estructura fija de la celda (ver docx_extract): p5 = "C.I:", p6 = "Aclaración: ____________________"
    for p in cell.paragraphs:
        if p.text.strip().startswith("C.I:"):
            _set_paragraph_text(p, f"C.I: {_v(ci)}")
        elif p.text.strip().startswith("Aclaración:"):
            _set_paragraph_text(p, f"Aclaración: {_v(nombre)}")


def generar_docx_entrega_equipo(a: dict) -> bytes:
    d = docx.Document(str(_TPL_ENTREGA_EQUIPO))

    # Tabla 1: Código de documento | Fecha de emisión
    _set_paragraph_text(d.tables[1].rows[0].cells[0].paragraphs[0], f"Código de documento: {_v(a.get('numero'))}")
    _set_paragraph_text(d.tables[1].rows[0].cells[1].paragraphs[0], f"Fecha de emisión: {_v(a.get('fecha_emision'))}")

    # Párrafo 3: oración de entrega — se reescribe entera en un solo run
    # (el original trae la oración partida en runs con datos de ejemplo
    # ya cargados, ej. "Roque"/"Esteche" en runs separados — no son
    # placeholders limpios para reemplazar uno por uno).
    entrega_nombre = _v(a.get("entrega_nombre")) or "____________________"
    entrega_cargo = a.get("entrega_cargo")
    recibe_nombre = _v(a.get("recibe_nombre")) or "____________________"
    recibe_ci = _v(a.get("recibe_ci")) or "____________________"
    recibe_cargo = a.get("recibe_cargo")
    cargo_entrega_txt = f", {entrega_cargo}" if entrega_cargo else ""
    cargo_recibe_txt = f", quien se desempeña como {recibe_cargo}" if recibe_cargo else ""
    nueva_oracion = (
        f"Por medio del presente documento, {entrega_nombre}{cargo_entrega_txt}, hace entrega a "
        f"{recibe_nombre}, con C.I. N° {recibe_ci}{cargo_recibe_txt}, del equipo informático que se "
        "detalla a continuación."
    )
    _set_paragraph_text(d.paragraphs[3], nueva_oracion)

    # Tabla 2: Insumo | Marca | Modelo | N° de Serie | Especificaciones
    _set_cell(d.tables[2], 1, 0, _v(a.get("insumo")))
    _set_cell(d.tables[2], 1, 1, _v(a.get("marca")))
    _set_cell(d.tables[2], 1, 2, _v(a.get("modelo")))
    _set_cell(d.tables[2], 1, 3, _v(a.get("nro_serie")))
    _set_cell(d.tables[2], 1, 4, _v(a.get("especificaciones")))

    # Tabla 3: Accesorios (checkboxes) — el símbolo "☐  " y la etiqueta
    # ("Cargador", "Funda"+" / Maletín", etc.) están en runs SEPARADOS
    # dentro del mismo párrafo, no en un único run — hay que concatenar
    # todos los runs del párrafo para leer la etiqueta completa.
    accesorios = set(a.get("accesorios") or [])
    for cell in d.tables[3].rows[0].cells:
        for p in cell.paragraphs:
            _check_symbol_in_paragraph(p, lambda label: label in accesorios)

    # "Otros: [Detallar accesorios adicionales, si corresponde]"
    if a.get("accesorios_otros"):
        p6 = d.paragraphs[6]
        if len(p6.runs) > 1:
            p6.runs[1].text = a["accesorios_otros"]

    # Tabla 4: Estado del equipo a la entrega (checkboxes Nuevo / Usado)
    estado = (a.get("estado_equipo") or "").strip()
    _check_estado_table(d.tables[4], estado)

    # "Observaciones: [Observaciones]" (entrega)
    p8 = d.paragraphs[8]
    if len(p8.runs) > 1:
        p8.runs[1].text = _v(a.get("observaciones_entrega")) or "Sin observaciones"

    # Tabla 5: Firma Entrega / Recibe (primera, sección 1)
    _fill_entrega_recibe_cell(d.tables[5].rows[0].cells[0], "Entrega", a.get("entrega_nombre"), a.get("entrega_ci"), a.get("entrega_cargo"))
    _fill_entrega_recibe_cell(d.tables[5].rows[0].cells[1], "Recibe", a.get("recibe_nombre"), a.get("recibe_ci"), a.get("recibe_cargo"))

    # Sección 2: devolución
    if a.get("fecha_devolucion"):
        _set_paragraph_text(d.paragraphs[13], f"Fecha: {_v(a.get('fecha_devolucion'))}")
    if a.get("motivo_devolucion"):
        _set_paragraph_text(d.paragraphs[14], f"Motivo de la devolución: {_v(a.get('motivo_devolucion'))}")

    estado_dev = (a.get("estado_equipo_devolucion") or "").strip()
    if estado_dev:
        _check_estado_table(d.tables[6], estado_dev)

    if a.get("observaciones_devolucion"):
        _set_paragraph_text(d.paragraphs[16], f"Observaciones: {_v(a.get('observaciones_devolucion'))}")

    if any([a.get("fecha_devolucion"), a.get("motivo_devolucion"), a.get("observaciones_devolucion")]):
        _fill_entrega_recibe_cell(d.tables[7].rows[0].cells[0], "Entrega", a.get("recibe_nombre"), a.get("recibe_ci"), a.get("recibe_cargo"))
        _fill_entrega_recibe_cell(d.tables[7].rows[0].cells[1], "Recibe", a.get("entrega_nombre"), a.get("entrega_ci"), a.get("entrega_cargo"))

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _check_estado_table(table, estado: str):
    def _matches(label: str) -> bool:
        return (label == "Nuevo" and estado == "Nuevo") or (label.startswith("Usado") and estado.startswith("Usado"))

    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            _check_symbol_in_paragraph(p, _matches)


def _check_symbol_in_paragraph(paragraph, matches_fn):
    """Busca el run que contiene "☐" en el párrafo, arma la etiqueta
    completa concatenando el resto de los runs del párrafo, y si
    `matches_fn(etiqueta)` da True reemplaza ese símbolo por "☑"."""
    symbol_run = None
    for r in paragraph.runs:
        if "☐" in r.text:
            symbol_run = r
            break
    if symbol_run is None:
        return
    label = "".join(r.text for r in paragraph.runs).replace("☐", "").strip()
    if matches_fn(label):
        symbol_run.text = symbol_run.text.replace("☐", "☑")


def _fill_entrega_recibe_cell(cell, rol: str, nombre, ci, cargo):
    for p in cell.paragraphs:
        t = p.text.strip()
        if t.startswith("Nombre:"):
            _set_paragraph_text(p, f"Nombre: {_v(nombre)}")
        elif t.startswith("C.I.:"):
            _set_paragraph_text(p, f"C.I.: {_v(ci)}")
        elif t.startswith("Cargo:"):
            _set_paragraph_text(p, f"Cargo: {_v(cargo)}")
